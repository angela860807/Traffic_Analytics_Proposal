from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from PIL import Image


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "산출물_단위테스트"
PROJECT = "번호판 인식 스마트 교통안전 시스템"
ORG = "에스트래픽\n스마트 모빌리티 DX Academy"
DATE = "2026. 06. 04"


def font(run, size=10, bold=False):
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(size)
    run.bold = bold


def p(doc, text="", align=None, size=10, bold=False):
    para = doc.add_paragraph()
    if align:
        para.alignment = align
    run = para.add_run(text)
    font(run, size, bold)
    return para


def h(doc, text, level=1):
    para = doc.add_heading(level=level)
    run = para.add_run(text)
    font(run, 14 if level == 1 else 12, True)


def shade(cell, color="D9EAF7"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, value in enumerate(headers):
        t.rows[0].cells[i].text = value
        shade(t.rows[0].cells[i])
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    for row in t.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    font(run, 8)
    doc.add_paragraph()


def cover(doc, title):
    for _ in range(4):
        doc.add_paragraph()
    p(doc, f"<{title}>", WD_ALIGN_PARAGRAPH.CENTER, 20, True)
    p(doc, f"<{PROJECT}>", WD_ALIGN_PARAGRAPH.CENTER, 15, True)
    for _ in range(7):
        doc.add_paragraph()
    p(doc, DATE, WD_ALIGN_PARAGRAPH.CENTER, 12)
    for line in ORG.splitlines():
        p(doc, line, WD_ALIGN_PARAGRAPH.CENTER, 12, line == "에스트래픽")
    doc.add_page_break()


def toc(doc, title, items):
    p(doc, title, WD_ALIGN_PARAGRAPH.CENTER, 13, True)
    p(doc, "- 목   차 -", WD_ALIGN_PARAGRAPH.CENTER, 12, True)
    doc.add_paragraph()
    for item in items:
        p(doc, item, size=10)
    doc.add_page_break()


def setup():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(2)
    sec.right_margin = Cm(2)
    doc.styles["Normal"].font.name = "Malgun Gothic"
    doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    doc.styles["Normal"].font.size = Pt(10)
    return doc


def add_image(doc, name, caption):
    path = ROOT / name
    if not path.exists():
        return
    p(doc, caption, size=9, bold=True)
    with Image.open(path) as img:
        w, hgt = img.size
    width = 15.5
    height = width * hgt / w
    if height > 10:
        height = 10
        width = height * w / hgt
    doc.add_picture(str(path), width=Cm(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()


def build_plan():
    doc = setup()
    cover(doc, "단위테스트 계획서")
    toc(
        doc,
        "단위테스트 계획서",
        [
            "1. 테스트 계획  3",
            "  1.1 테스트 목적  3",
            "  1.2 테스트 범위  3",
            "  1.3 테스트 일정  3",
            "  1.4 역할 및 책임  3",
            "2. 테스트 전략  4",
            "  2.1 테스트 유형  4",
            "  2.2 테스트 환경  4",
            "  2.3 테스트 도구  4",
            "3. 테스트 실행  5",
            "  3.1 테스트 항목  5",
            "  3.2 테스트 시나리오  5",
        ],
    )
    h(doc, "1. 테스트 계획")
    h(doc, "1.1 테스트 목적", 2)
    for item in [
        "FastAPI, Spring Boot, Vue 프론트엔드 주요 기능 정확성 검증",
        "API 응답, DB 저장, 화면 표시 결과의 정상 동작 확인",
        "Docker Compose 기반 최종 실행 환경에서 초기 결함 발견",
        "테스트 결과와 캡처 증빙을 통한 산출물 근거 확보",
    ]:
        p(doc, f"- {item}")
    h(doc, "1.2 테스트 범위", 2)
    for item in [
        "FastAPI 기능: Health Check, 이미지/영상 분석, 번호판/OCR 처리, Spring Boot 전송",
        "Spring Boot 기능: 감지 로그 저장, 분석 결과 저장, 과속 위반 조회",
        "프론트엔드 기능: 메인 화면, 감지 로그 화면, 분석 결과 표시",
        "DB 기능: PostgreSQL 감지 로그/분석 결과/과속 위반 데이터 조회",
        "Docker 기능: 전체 서비스 빌드 및 실행",
    ]:
        p(doc, f"- {item}")
    h(doc, "1.3 테스트 일정", 2)
    table(
        doc,
        ["단계", "일정", "담당자", "비고"],
        [
            ["테스트 준비", "2026-06-04", "QA/통합 담당", "Docker 및 의존성 확인"],
            ["단위테스트 실행", "2026-06-04", "각 파트 담당", "pytest, Gradle, Vitest"],
            ["결과 확인", "2026-06-04", "QA/통합 담당", "API/DB/화면 캡처"],
            ["결과서 작성", "2026-06-04", "QA/통합 담당", "증빙 이미지 정리"],
        ],
    )
    h(doc, "1.4 역할 및 책임", 2)
    for item in [
        "프로젝트 매니저: 테스트 범위 확인 및 최종 산출물 검토",
        "FastAPI 담당: AI 분석 API 및 Spring Boot 전송 기능 검증",
        "Spring Boot 담당: 저장 API, 조회 API, DB 연동 검증",
        "프론트엔드 담당: 화면 표시, API 응답 렌더링 검증",
        "QA/통합 담당: 전체 실행 환경, 테스트 수행, 결과서 작성",
    ]:
        p(doc, f"- {item}")
    h(doc, "2. 테스트 전략")
    h(doc, "2.1 테스트 유형", 2)
    p(doc, "- 기능 테스트")
    p(doc, "- API 응답 테스트")
    p(doc, "- DB 저장 확인 테스트")
    p(doc, "- 화면 표시 테스트")
    h(doc, "2.2 테스트 환경", 2)
    for item in [
        "Windows",
        "Docker Compose",
        "PostgreSQL 16",
        "Spring Boot 3.5.14 / Java 21",
        "FastAPI / Python 3.12",
        "Vue 3 / Vite 5",
        "Node v24.13.1 / npm 11.8.0",
    ]:
        p(doc, f"- {item}")
    h(doc, "2.3 테스트 도구", 2)
    for item in ["pytest", "Gradle/JUnit", "Vitest", "Vite build", "Swagger UI", "psql", "Docker Compose"]:
        p(doc, f"- {item}")
    h(doc, "3. 테스트 실행")
    h(doc, "3.1 테스트 항목", 2)
    for item in [
        "FastAPI 자동화 단위테스트",
        "Spring Boot 자동화 테스트",
        "프론트엔드 테스트 및 빌드",
        "FastAPI/Spring Boot API 응답 확인",
        "PostgreSQL 저장 데이터 조회",
        "프론트엔드 화면 표시 확인",
    ]:
        p(doc, f"- {item}")
    h(doc, "3.2 테스트 시나리오", 2)
    table(
        doc,
        ["테스트 ID", "항목", "시나리오", "기대 결과"],
        [
            ["UT-1.1", "FastAPI", "pytest 자동화 테스트 실행", "전체 테스트 통과"],
            ["UT-1.2", "Spring Boot", "Gradle test 실행", "BUILD SUCCESSFUL"],
            ["UT-1.3", "Frontend", "npm test 및 build 실행", "테스트 결과 확인 및 build 성공"],
            ["UT-1.4", "API", "Health/감지로그/과속위반 API 조회", "정상 응답"],
            ["UT-1.5", "DB", "PostgreSQL 테이블 조회", "저장 데이터 확인"],
            ["UT-1.6", "화면", "프론트 화면에서 결과 표시 확인", "화면 정상 출력"],
        ],
    )
    doc.save(OUT / "[KDT] 인천부평_1기_1조_단위테스트_계획서_양식맞춤_최종.docx")


def build_result():
    doc = setup()
    cover(doc, "단위테스트 결과서")
    toc(
        doc,
        "단위테스트 결과서",
        [
            "1. 테스트 항목  3",
            "  1.1 FastAPI 테스트  3",
            "  1.2 Spring Boot 테스트  4",
            "  1.3 프론트엔드 테스트  4",
            "  1.4 API/DB/화면 확인  5",
            "2. 테스트 결과  6",
            "  2.1 정상/오류 확인  6",
            "3. 개선사항  7",
            "  3.1 결함 및 보완사항  7",
            "4. 증빙자료  8",
        ],
    )
    h(doc, "1. 테스트 항목")
    h(doc, "1.1 FastAPI 테스트", 2)
    p(doc, "- pytest 자동화 테스트 실행")
    p(doc, "- Health Check API 응답 확인")
    add_image(doc, "UT_FAST_001_pytest_result.png", "FastAPI pytest 결과")
    add_image(doc, "API_001_fastapi_health.png", "FastAPI Health API 응답")
    h(doc, "1.2 Spring Boot 테스트", 2)
    p(doc, "- Gradle/JUnit 테스트 실행")
    p(doc, "- 감지 로그 및 과속 위반 API 응답 확인")
    add_image(doc, "UT_SPR_001_gradle_test_result.png", "Spring Boot Gradle test 결과")
    add_image(doc, "API_002_detection_logs_response.png", "감지 로그 API 응답")
    add_image(doc, "API_003_speed_violations_response.png", "과속 위반 API 응답")
    h(doc, "1.3 프론트엔드 테스트", 2)
    p(doc, "- Vitest 테스트 실행")
    p(doc, "- 메인 화면 및 감지 로그 화면 출력 확인")
    add_image(doc, "UT_FE_001_npm_test_result.png", "Frontend Vitest 결과")
    add_image(doc, "FE_001_main_page.png", "프론트엔드 메인 화면")
    add_image(doc, "FE_003_detection_logs.png", "프론트엔드 감지 로그 화면")
    h(doc, "1.4 API/DB/화면 확인", 2)
    p(doc, "- FastAPI 분석 결과의 Spring Boot 저장 확인")
    p(doc, "- PostgreSQL 테이블별 저장 데이터 조회")
    add_image(doc, "IT_003_front_to_fastapi_detection.png", "프론트엔드에서 FastAPI 분석 결과 확인")
    add_image(doc, "IT_004_fastapi_to_spring_saved.png", "FastAPI 분석 결과 Spring Boot 저장 확인")
    add_image(doc, "IT_005_01_detection_logs_query.png", "detection_logs DB 조회")
    add_image(doc, "IT_005_02_detection_analysis_results_query.png", "detection_analysis_results DB 조회")
    add_image(doc, "IT_005_03_speed_violations_query.png", "speed_violations DB 조회")
    h(doc, "2. 테스트 결과")
    h(doc, "2.1 정상/오류 확인", 2)
    table(
        doc,
        ["테스트 ID", "항목", "시나리오", "결과"],
        [
            ["UT-1.1", "FastAPI pytest", "FastAPI 단위테스트 64건 실행", "정상"],
            ["UT-1.2", "FastAPI Health", "/health API 응답 확인", "정상"],
            ["UT-1.3", "Spring Boot", "Gradle test 실행", "정상"],
            ["UT-1.4", "Frontend Vitest", "프론트엔드 테스트 실행", "오류"],
            ["UT-1.5", "Frontend 화면", "메인/감지 로그 화면 확인", "정상"],
            ["UT-1.6", "API 조회", "감지 로그/과속 위반 API 응답 확인", "정상"],
            ["UT-1.7", "DB 조회", "PostgreSQL 저장 데이터 조회", "정상"],
            ["UT-1.8", "연동 확인", "FastAPI 결과 Spring Boot 저장 확인", "정상"],
            ["완료율", "-", "정상 7건, 오류 1건", "87.5%"],
        ],
    )
    h(doc, "3. 개선사항")
    h(doc, "3.1 결함 및 보완사항", 2)
    p(doc, "- FastAPI: PaddleOCR ocr() 호출 관련 deprecation warning 발생, 향후 predict() 방식 전환 검토")
    p(doc, "- FastAPI: ccache 미설치 warning 발생, 기능 오류는 아니나 컴파일 캐시 도구 설치 검토")
    p(doc, "- Frontend: DataState.vue import 경로 또는 파일 누락으로 Vitest suite 실패, 테스트 코드와 실제 컴포넌트 구조 동기화 필요")
    p(doc, "- Frontend: Vite build 시 chunk size warning 발생, 추후 dynamic import/manualChunks 적용 검토")
    h(doc, "4. 증빙자료")
    add_image(doc, "IT_001_docker_compose_build.png", "Docker Compose 빌드 및 실행 확인")
    doc.save(OUT / "[KDT] 인천부평_1기_1조_단위테스트_결과서_양식맞춤_최종.docx")


def main():
    OUT.mkdir(exist_ok=True)
    build_plan()
    build_result()
    print(OUT)


if __name__ == "__main__":
    main()
