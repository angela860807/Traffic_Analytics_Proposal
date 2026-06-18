from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from PIL import Image


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "산출물_통합테스트"
PROJECT = "번호판 인식 스마트 교통안전 시스템"
ORG = "에스트래픽\n스마트 모빌리티 DX Academy"
DATE = "2026. 06. 04"


def font(run, size=10, bold=False):
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(size)
    run.bold = bold


def para(doc, text="", align=None, size=10, bold=False):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    r = p.add_run(text)
    font(r, size, bold)
    return p


def heading(doc, text, level=1):
    p = doc.add_heading(level=level)
    r = p.add_run(text)
    font(r, 14 if level == 1 else 12, True)


def shade(cell, color="D9EAF7"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, value in enumerate(headers):
        t.rows[0].cells[i].text = value
        shade(t.rows[0].cells[i])
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    font(r, 8)
    doc.add_paragraph()


def setup():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(2)
    sec.right_margin = Cm(2)
    doc.styles["Normal"].font.name = "Malgun Gothic"
    doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    doc.styles["Normal"].font.size = Pt(10)
    return doc


def cover(doc, title):
    for _ in range(4):
        doc.add_paragraph()
    para(doc, f"<{title}>", WD_ALIGN_PARAGRAPH.CENTER, 20, True)
    para(doc, f"<{PROJECT}>", WD_ALIGN_PARAGRAPH.CENTER, 15, True)
    for _ in range(7):
        doc.add_paragraph()
    para(doc, DATE, WD_ALIGN_PARAGRAPH.CENTER, 12)
    for line in ORG.splitlines():
        para(doc, line, WD_ALIGN_PARAGRAPH.CENTER, 12, line == "에스트래픽")
    doc.add_page_break()


def toc(doc, title, items):
    para(doc, title, WD_ALIGN_PARAGRAPH.CENTER, 13, True)
    para(doc, "- 목   차 -", WD_ALIGN_PARAGRAPH.CENTER, 12, True)
    doc.add_paragraph()
    for item in items:
        para(doc, item, size=10)
    doc.add_page_break()


def add_image(doc, name, caption):
    path = ROOT / name
    if not path.exists():
        para(doc, f"[증빙 누락] {name} - {caption}", size=9, bold=True)
        return
    para(doc, caption, size=9, bold=True)
    with Image.open(path) as img:
        w, h = img.size
    width = 15.5
    height = width * h / w
    if height > 10:
        height = 10
        width = height * w / h
    doc.add_picture(str(path), width=Cm(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()


PLAN_MD = """# 통합테스트 계획서

## 1. 테스트 계획

### 1.1 테스트 목적
- Docker Compose 기반 최종 실행 환경에서 Frontend, FastAPI, Spring Boot, PostgreSQL 간 연동 상태를 검증한다.
- 프론트엔드 분석 요청부터 FastAPI 분석 처리, Spring Boot 저장 API 호출, PostgreSQL 저장, 프론트엔드 조회 화면 표시까지 전체 흐름을 확인한다.
- 발표 및 시연 환경에서 실제 사용자 흐름이 정상 동작하는지 확인한다.
- 각 모듈이 단독으로 정상 동작하는 것을 넘어, API 계약과 데이터 필드가 모듈 사이에서 누락 없이 전달되는지 확인한다.
- 장애 발생 시 어느 구간에서 문제가 발생했는지 추적할 수 있도록 단계별 관찰 지점을 정의한다.

### 1.2 테스트 범위
- Docker Compose 전체 서비스 실행
- Frontend에서 FastAPI 분석 기능 호출
- FastAPI 분석 결과의 Spring Boot 저장 API 전송
- Spring Boot의 PostgreSQL 저장 및 조회
- Spring Boot 조회 API 응답의 프론트엔드 표시

### 1.3 테스트 일정
| 단계 | 일정 | 담당자 | 비고 |
|---|---|---|---|
| 통합 환경 준비 | 2026-06-04 | QA/통합 담당 | Docker Compose 실행 |
| 통합 시나리오 수행 | 2026-06-04 | QA/통합 담당 | API/DB/화면 확인 |
| 결과서 작성 | 2026-06-04 | QA/통합 담당 | 캡처 증빙 정리 |

### 1.4 역할 및 책임
- FastAPI 담당: 분석 API 및 Spring Boot 전송 검증
- Spring Boot 담당: 저장 API, 조회 API, DB 저장 검증
- Frontend 담당: 분석 요청 화면 및 결과 조회 화면 검증
- QA/통합 담당: 전체 흐름 수행, 캡처 증빙 수집, 결과서 작성

### 1.5 통합 대상 구성
| 구간 | 송신 모듈 | 수신 모듈 | 주요 데이터 | 검증 포인트 |
|---|---|---|---|---|
| 1 | Frontend | FastAPI | cameraCode, capturedAt, image/frame | 분석 요청 정상 수신 및 결과 응답 |
| 2 | FastAPI | Spring Boot | plateNumber, detectionType, confidenceScore, imagePath, detectedAt | 저장 API 응답 및 상태값 확인 |
| 3 | Spring Boot | PostgreSQL | detection_logs, detection_analysis_results, speed_violations | 테이블 row 생성 및 최신 데이터 조회 |
| 4 | Frontend | Spring Boot | 감지 로그/과속 위반 조회 요청 | API 응답 기반 화면 표시 |

## 2. 테스트 전략
- End-to-End 흐름 테스트
- API 연동 테스트
- DB 저장 확인 테스트
- 화면 표시 확인 테스트
- Docker 실행 환경 확인

### 2.1 선행 조건
- Docker Desktop 실행 상태
- PostgreSQL, Spring Boot, FastAPI, Frontend 컨테이너 실행 가능 상태
- Spring Boot 내부 API Key와 FastAPI 환경변수의 저장 API 경로 일치
- 프론트엔드 API base URL이 Docker/nginx 프록시 기준으로 요청 가능한 상태
- 테스트용 이미지 또는 mock 분석 요청 데이터 준비

### 2.2 종료 조건
- 전체 컨테이너가 정상 실행된다.
- 프론트엔드에서 FastAPI 분석 요청이 수행된다.
- FastAPI 분석 결과가 Spring Boot 저장 API로 전달된다.
- PostgreSQL에서 감지 로그/분석 결과/과속 위반 데이터가 조회된다.
- Spring Boot 조회 API 응답이 확인된다.
- 프론트엔드 화면에서 저장/조회 결과가 표시된다.

### 2.3 실패 판단 기준
- 컨테이너 빌드 또는 실행 실패
- FastAPI 분석 API 4xx/5xx 응답
- FastAPI to Spring Boot 저장 요청 실패 또는 API Key 오류
- DB 테이블에 저장 row 미생성
- Spring Boot 조회 API 응답 실패
- 프론트엔드 화면에서 결과 표시 누락

## 3. 테스트 실행
| 테스트 ID | 항목 | 시나리오 | 기대 결과 |
|---|---|---|---|
| IT-1.1 | Docker 전체 실행 | docker compose up --build 수행 | 전체 서비스 정상 실행 |
| IT-1.2 | Frontend to FastAPI | 프론트에서 분석 요청 수행 | FastAPI 분석 결과 응답 |
| IT-1.3 | FastAPI to Spring Boot | 분석 결과 저장 API 전송 | Spring Boot 저장 응답 반환 |
| IT-1.4 | Spring Boot to PostgreSQL | 감지 로그/분석 결과/과속 위반 저장 확인 | DB row 조회 가능 |
| IT-1.5 | Spring Boot API 조회 | 감지 로그/과속 위반 조회 API 호출 | 정상 응답 |
| IT-1.6 | Frontend 결과 표시 | 저장/조회 결과를 프론트 화면에서 확인 | 화면 정상 표시 |

### 3.1 상세 통합 시나리오
| 단계 | 수행 내용 | 확인 데이터 | 증빙 |
|---|---|---|---|
| 1 | Docker Compose로 전체 서비스 실행 | 컨테이너 빌드/기동 로그 | IT_001 |
| 2 | 프론트엔드에서 분석 요청 수행 | FastAPI 분석 결과 표시 | IT_003 |
| 3 | FastAPI가 Spring Boot 저장 API 호출 | 저장 응답, logId/상태값 | IT_004 |
| 4 | PostgreSQL 감지 로그 조회 | detection_logs 최신 row | IT_005_01 |
| 5 | PostgreSQL 분석 결과 조회 | detection_analysis_results 최신 row | IT_005_02 |
| 6 | PostgreSQL 과속 위반 조회 | speed_violations 조회 결과 | IT_005_03 |
| 7 | Spring Boot 조회 API 확인 | 감지 로그/과속 위반 API 응답 | API_002, API_003 |
| 8 | 프론트엔드 화면 확인 | 감지 로그 화면 표시 | FE_003 |
"""


RESULT_MD = """# 통합테스트 결과서

## 1. 테스트 항목
- Docker 전체 실행
- Frontend to FastAPI 분석 요청
- FastAPI to Spring Boot 저장
- Spring Boot to PostgreSQL 저장
- Spring Boot 조회 API 응답
- Frontend 결과 화면 표시

## 1.1 통합 흐름 요약
본 통합테스트는 단일 기능의 정상 여부가 아니라, 사용자가 프론트엔드에서 분석을 요청했을 때 분석 결과가 백엔드와 DB를 거쳐 다시 조회 화면에 표시되는 전체 흐름을 검증하였다.
검증 흐름은 `Frontend → FastAPI → Spring Boot → PostgreSQL → Spring Boot API → Frontend` 순서로 수행하였다.

## 1.2 주요 인터페이스 확인 항목
| 인터페이스 | 확인 항목 |
|---|---|
| Frontend to FastAPI | 분석 요청 전송, 분석 결과 응답 표시 |
| FastAPI to Spring Boot | 내부 API Key, 저장 API 경로, 분석 결과 payload |
| Spring Boot to PostgreSQL | 감지 로그, 분석 결과, 과속 위반 테이블 저장 |
| Spring Boot to Frontend | 조회 API 응답 및 화면 표시 |

## 2. 테스트 결과
| 테스트 ID | 항목 | 시나리오 | 결과 |
|---|---|---|---|
| IT-1.1 | Docker 전체 실행 | docker compose up --build 수행 | 정상 |
| IT-1.2 | Frontend to FastAPI | 프론트에서 분석 요청 후 결과 확인 | 정상 |
| IT-1.3 | FastAPI to Spring Boot | FastAPI 분석 결과를 Spring Boot 저장 API로 전송 | 정상 |
| IT-1.4 | PostgreSQL 저장 | detection_logs, detection_analysis_results, speed_violations 조회 | 정상 |
| IT-1.5 | Spring Boot API 조회 | 감지 로그/과속 위반 조회 API 응답 확인 | 정상 |
| IT-1.6 | Frontend 결과 표시 | 감지 로그 화면에서 결과 표시 확인 | 정상 |
| 완료율 | - | 정상 6건, 오류 0건 | 100% |

## 2.1 단계별 확인 결과
| 단계 | 검증 내용 | 확인 결과 |
|---|---|---|
| Docker 실행 | 전체 서비스 빌드 및 실행 로그 확인 | 정상 |
| 분석 요청 | 프론트엔드에서 FastAPI 분석 결과 화면 확인 | 정상 |
| 저장 요청 | FastAPI가 Spring Boot 저장 API 응답 수신 | 정상 |
| DB 저장 | detection_logs, detection_analysis_results, speed_violations 조회 | 정상 |
| API 조회 | Spring Boot 감지 로그/과속 위반 조회 API 응답 | 정상 |
| 화면 표시 | 프론트엔드 감지 로그 화면 표시 | 정상 |

## 2.2 통합 관점 검토
- Docker Compose 기반으로 실행 환경을 통일하여 로컬 실행 편차를 줄였다.
- FastAPI 분석 결과가 Spring Boot 저장 API로 전달되는 경로를 확인하였다.
- 저장 API 응답뿐 아니라 PostgreSQL 테이블 조회를 통해 실제 저장 여부를 확인하였다.
- DB에 저장된 데이터가 Spring Boot 조회 API와 프론트엔드 화면에서 확인되는지 검증하였다.
- 통합 흐름 내 주요 실패 지점인 컨테이너 실행, API 응답, DB 저장, 화면 표시를 단계별로 분리하여 확인하였다.

## 3. 개선사항
- Docker 실행 및 API/DB 조회 절차를 README에 고정하여 재현성을 높인다.
- 통합 시나리오를 자동화 테스트로 확장하여 수동 캡처 의존도를 줄인다.
- 프론트엔드 테스트 코드와 실제 컴포넌트 구조를 동기화하여 단위테스트 실패 항목을 보완한다.
- FastAPI to Spring Boot 저장 실패 시 재시도 횟수, 오류 메시지, 실패 payload를 별도 로그로 남기는 방식을 강화한다.
- 프론트엔드에서 API 장애 상태를 사용자에게 명확히 표시할 수 있도록 error/empty/loading 상태 검증을 통합 시나리오에 추가한다.
"""


def write_md():
    OUT.mkdir(exist_ok=True)
    (OUT / "[KDT] 인천부평_1기_1조_통합테스트_계획서_양식맞춤_최종.md").write_text(
        PLAN_MD.strip() + "\n", encoding="utf-8"
    )
    (OUT / "[KDT] 인천부평_1기_1조_통합테스트_결과서_양식맞춤_최종.md").write_text(
        RESULT_MD.strip() + "\n", encoding="utf-8"
    )


def build_plan():
    doc = setup()
    cover(doc, "통합테스트 계획서")
    toc(
        doc,
        "통합테스트 계획서",
        [
            "1. 테스트 계획  3",
            "  1.1 테스트 목적  3",
            "  1.2 테스트 범위  3",
            "  1.3 테스트 일정  3",
            "  1.4 역할 및 책임  3",
            "  1.5 통합 대상 구성  4",
            "2. 테스트 전략  5",
            "  2.1 테스트 유형  5",
            "  2.2 테스트 환경  5",
            "  2.3 테스트 도구  5",
            "  2.4 선행/종료 조건  6",
            "3. 테스트 실행  7",
            "  3.1 테스트 항목  7",
            "  3.2 테스트 시나리오  7",
            "  3.3 상세 통합 시나리오  8",
        ],
    )
    heading(doc, "1. 테스트 계획")
    heading(doc, "1.1 테스트 목적", 2)
    for item in [
        "Docker Compose 기반 최종 실행 환경에서 Frontend, FastAPI, Spring Boot, PostgreSQL 간 연동 상태 검증",
        "프론트엔드 분석 요청부터 DB 저장 및 화면 조회까지 전체 흐름 확인",
        "발표 및 시연 환경에서 실제 사용자 흐름이 정상 동작하는지 확인",
        "API 계약과 주요 데이터 필드가 모듈 사이에서 누락 없이 전달되는지 확인",
        "장애 발생 시 문제 구간을 추적할 수 있도록 단계별 관찰 지점 정의",
    ]:
        para(doc, f"- {item}")
    heading(doc, "1.2 테스트 범위", 2)
    for item in [
        "Docker Compose 전체 서비스 실행",
        "Frontend에서 FastAPI 분석 기능 호출",
        "FastAPI 분석 결과의 Spring Boot 저장 API 전송",
        "Spring Boot의 PostgreSQL 저장 및 조회",
        "Spring Boot 조회 API 응답의 프론트엔드 표시",
    ]:
        para(doc, f"- {item}")
    heading(doc, "1.3 테스트 일정", 2)
    table(
        doc,
        ["단계", "일정", "담당자", "비고"],
        [
            ["통합 환경 준비", "2026-06-04", "QA/통합 담당", "Docker Compose 실행"],
            ["통합 시나리오 수행", "2026-06-04", "QA/통합 담당", "API/DB/화면 확인"],
            ["결과서 작성", "2026-06-04", "QA/통합 담당", "캡처 증빙 정리"],
        ],
    )
    heading(doc, "1.4 역할 및 책임", 2)
    for item in [
        "FastAPI 담당: 분석 API 및 Spring Boot 전송 검증",
        "Spring Boot 담당: 저장 API, 조회 API, DB 저장 검증",
        "Frontend 담당: 분석 요청 화면 및 결과 조회 화면 검증",
        "QA/통합 담당: 전체 흐름 수행, 캡처 증빙 수집, 결과서 작성",
    ]:
        para(doc, f"- {item}")
    heading(doc, "1.5 통합 대상 구성", 2)
    table(
        doc,
        ["구간", "송신 모듈", "수신 모듈", "주요 데이터", "검증 포인트"],
        [
            ["1", "Frontend", "FastAPI", "cameraCode, capturedAt, image/frame", "분석 요청 정상 수신 및 결과 응답"],
            ["2", "FastAPI", "Spring Boot", "plateNumber, detectionType, confidenceScore, imagePath, detectedAt", "저장 API 응답 및 상태값 확인"],
            ["3", "Spring Boot", "PostgreSQL", "detection_logs, detection_analysis_results, speed_violations", "테이블 row 생성 및 최신 데이터 조회"],
            ["4", "Frontend", "Spring Boot", "감지 로그/과속 위반 조회 요청", "API 응답 기반 화면 표시"],
        ],
    )
    heading(doc, "2. 테스트 전략")
    heading(doc, "2.1 테스트 유형", 2)
    for item in ["End-to-End 흐름 테스트", "API 연동 테스트", "DB 저장 확인 테스트", "화면 표시 확인 테스트", "Docker 실행 환경 확인"]:
        para(doc, f"- {item}")
    heading(doc, "2.2 테스트 환경", 2)
    for item in ["Windows", "Docker Compose", "PostgreSQL 16", "Spring Boot 3.5.14", "FastAPI / Python 3.12", "Vue 3 / Vite 5"]:
        para(doc, f"- {item}")
    heading(doc, "2.3 테스트 도구", 2)
    for item in ["Docker Compose", "Swagger UI", "PowerShell/curl", "PostgreSQL psql", "브라우저"]:
        para(doc, f"- {item}")
    heading(doc, "2.4 선행/종료 조건", 2)
    table(
        doc,
        ["구분", "내용"],
        [
            ["선행 조건", "Docker Desktop 실행, 전체 컨테이너 실행 가능, FastAPI/Spring Boot API 경로 및 내부 API Key 설정"],
            ["선행 조건", "프론트엔드 API 요청 경로가 Docker/nginx 또는 localhost 기준으로 접근 가능한 상태"],
            ["종료 조건", "분석 요청, 저장 API 응답, DB 조회, API 조회, 프론트 화면 표시가 모두 정상"],
            ["실패 기준", "컨테이너 실행 실패, API 4xx/5xx, DB 저장 누락, 프론트 화면 표시 실패"],
        ],
    )
    heading(doc, "3. 테스트 실행")
    heading(doc, "3.1 테스트 항목", 2)
    for item in [
        "Docker 전체 서비스 실행",
        "Frontend to FastAPI 분석 요청",
        "FastAPI to Spring Boot 저장",
        "Spring Boot to PostgreSQL 저장",
        "Spring Boot 조회 API 응답",
        "Frontend 결과 화면 표시",
    ]:
        para(doc, f"- {item}")
    heading(doc, "3.2 테스트 시나리오", 2)
    table(
        doc,
        ["테스트 ID", "항목", "시나리오", "기대 결과"],
        [
            ["IT-1.1", "Docker 전체 실행", "docker compose up --build 수행", "전체 서비스 정상 실행"],
            ["IT-1.2", "Frontend to FastAPI", "프론트에서 분석 요청 수행", "FastAPI 분석 결과 응답"],
            ["IT-1.3", "FastAPI to Spring Boot", "분석 결과 저장 API 전송", "Spring Boot 저장 응답 반환"],
            ["IT-1.4", "Spring Boot to PostgreSQL", "감지 로그/분석 결과/과속 위반 저장 확인", "DB row 조회 가능"],
            ["IT-1.5", "Spring Boot API 조회", "감지 로그/과속 위반 조회 API 호출", "정상 응답"],
            ["IT-1.6", "Frontend 결과 표시", "저장/조회 결과를 프론트 화면에서 확인", "화면 정상 표시"],
        ],
    )
    heading(doc, "3.3 상세 통합 시나리오", 2)
    table(
        doc,
        ["단계", "수행 내용", "확인 데이터", "증빙"],
        [
            ["1", "Docker Compose로 전체 서비스 실행", "컨테이너 빌드/기동 로그", "IT_001"],
            ["2", "프론트엔드에서 분석 요청 수행", "FastAPI 분석 결과 표시", "IT_003"],
            ["3", "FastAPI가 Spring Boot 저장 API 호출", "저장 응답, logId/상태값", "IT_004"],
            ["4", "PostgreSQL 감지 로그 조회", "detection_logs 최신 row", "IT_005_01"],
            ["5", "PostgreSQL 분석 결과 조회", "detection_analysis_results 최신 row", "IT_005_02"],
            ["6", "PostgreSQL 과속 위반 조회", "speed_violations 조회 결과", "IT_005_03"],
            ["7", "Spring Boot 조회 API 확인", "감지 로그/과속 위반 API 응답", "API_002, API_003"],
            ["8", "프론트엔드 화면 확인", "감지 로그 화면 표시", "FE_003"],
        ],
    )
    doc.save(OUT / "[KDT] 인천부평_1기_1조_통합테스트_계획서_양식맞춤_최종.docx")


def build_result():
    doc = setup()
    cover(doc, "통합테스트 결과서")
    toc(
        doc,
        "통합테스트 결과서",
        [
            "1. 테스트 항목  3",
            "  1.1 Docker 전체 실행  3",
            "  1.2 Frontend to FastAPI  4",
            "  1.3 FastAPI to Spring Boot  5",
            "  1.4 Spring Boot to PostgreSQL  6",
            "  1.5 Spring Boot API 및 화면 조회  7",
            "  1.6 통합 흐름 요약  8",
            "2. 테스트 결과  9",
            "  2.1 정상/오류 확인  9",
            "  2.2 단계별 확인 결과  10",
            "  2.3 통합 관점 검토  10",
            "3. 개선사항  11",
            "  3.1 결함 및 보완사항  11",
        ],
    )
    heading(doc, "1. 테스트 항목")
    heading(doc, "1.1 Docker 전체 실행", 2)
    para(doc, "- docker compose up --build를 통해 PostgreSQL, Spring Boot, FastAPI, Frontend 실행 확인")
    add_image(doc, "IT_001_docker_compose_build.png", "Docker Compose 빌드 및 실행 확인")
    heading(doc, "1.2 Frontend to FastAPI", 2)
    para(doc, "- 프론트엔드에서 분석 요청 후 FastAPI 분석 결과 표시 확인")
    add_image(doc, "IT_003_front_to_fastapi_detection.png", "프론트엔드에서 FastAPI 분석 결과 확인")
    heading(doc, "1.3 FastAPI to Spring Boot", 2)
    para(doc, "- FastAPI 분석 결과를 Spring Boot 저장 API로 전송 후 응답 확인")
    add_image(doc, "IT_004_fastapi_to_spring_saved.png", "FastAPI 분석 결과 Spring Boot 저장 응답")
    heading(doc, "1.4 Spring Boot to PostgreSQL", 2)
    para(doc, "- Spring Boot 저장 결과가 PostgreSQL 테이블에 반영되었는지 조회")
    add_image(doc, "IT_005_01_detection_logs_query.png", "detection_logs DB 조회")
    add_image(doc, "IT_005_02_detection_analysis_results_query.png", "detection_analysis_results DB 조회")
    add_image(doc, "IT_005_03_speed_violations_query.png", "speed_violations DB 조회")
    heading(doc, "1.5 Spring Boot API 및 화면 조회", 2)
    para(doc, "- Spring Boot 조회 API 응답 및 프론트엔드 감지 로그 화면 표시 확인")
    add_image(doc, "API_002_detection_logs_response.png", "감지 로그 조회 API 응답")
    add_image(doc, "API_003_speed_violations_response.png", "과속 위반 조회 API 응답")
    add_image(doc, "FE_003_detection_logs.png", "프론트엔드 감지 로그 화면")
    heading(doc, "1.6 통합 흐름 요약", 2)
    para(doc, "본 통합테스트는 단일 기능의 정상 여부가 아니라, 사용자가 프론트엔드에서 분석을 요청했을 때 분석 결과가 백엔드와 DB를 거쳐 다시 조회 화면에 표시되는 전체 흐름을 검증하였다.")
    para(doc, "검증 흐름은 Frontend -> FastAPI -> Spring Boot -> PostgreSQL -> Spring Boot API -> Frontend 순서로 수행하였다.")
    table(
        doc,
        ["인터페이스", "확인 항목"],
        [
            ["Frontend to FastAPI", "분석 요청 전송, 분석 결과 응답 표시"],
            ["FastAPI to Spring Boot", "내부 API Key, 저장 API 경로, 분석 결과 payload"],
            ["Spring Boot to PostgreSQL", "감지 로그, 분석 결과, 과속 위반 테이블 저장"],
            ["Spring Boot to Frontend", "조회 API 응답 및 화면 표시"],
        ],
    )
    heading(doc, "2. 테스트 결과")
    heading(doc, "2.1 정상/오류 확인", 2)
    table(
        doc,
        ["테스트 ID", "항목", "시나리오", "결과"],
        [
            ["IT-1.1", "Docker 전체 실행", "docker compose up --build 수행", "정상"],
            ["IT-1.2", "Frontend to FastAPI", "프론트에서 분석 요청 후 결과 확인", "정상"],
            ["IT-1.3", "FastAPI to Spring Boot", "FastAPI 분석 결과를 Spring Boot 저장 API로 전송", "정상"],
            ["IT-1.4", "PostgreSQL 저장", "detection_logs, detection_analysis_results, speed_violations 조회", "정상"],
            ["IT-1.5", "Spring Boot API 조회", "감지 로그/과속 위반 조회 API 응답 확인", "정상"],
            ["IT-1.6", "Frontend 결과 표시", "감지 로그 화면에서 결과 표시 확인", "정상"],
            ["완료율", "-", "정상 6건, 오류 0건", "100%"],
        ],
    )
    heading(doc, "2.2 단계별 확인 결과", 2)
    table(
        doc,
        ["단계", "검증 내용", "확인 결과"],
        [
            ["Docker 실행", "전체 서비스 빌드 및 실행 로그 확인", "정상"],
            ["분석 요청", "프론트엔드에서 FastAPI 분석 결과 화면 확인", "정상"],
            ["저장 요청", "FastAPI가 Spring Boot 저장 API 응답 수신", "정상"],
            ["DB 저장", "detection_logs, detection_analysis_results, speed_violations 조회", "정상"],
            ["API 조회", "Spring Boot 감지 로그/과속 위반 조회 API 응답", "정상"],
            ["화면 표시", "프론트엔드 감지 로그 화면 표시", "정상"],
        ],
    )
    heading(doc, "2.3 통합 관점 검토", 2)
    for item in [
        "Docker Compose 기반으로 실행 환경을 통일하여 로컬 실행 편차를 줄였다.",
        "FastAPI 분석 결과가 Spring Boot 저장 API로 전달되는 경로를 확인하였다.",
        "저장 API 응답뿐 아니라 PostgreSQL 테이블 조회를 통해 실제 저장 여부를 확인하였다.",
        "DB에 저장된 데이터가 Spring Boot 조회 API와 프론트엔드 화면에서 확인되는지 검증하였다.",
        "통합 흐름 내 주요 실패 지점인 컨테이너 실행, API 응답, DB 저장, 화면 표시를 단계별로 분리하여 확인하였다.",
    ]:
        para(doc, f"- {item}")
    heading(doc, "3. 개선사항")
    heading(doc, "3.1 결함 및 보완사항", 2)
    for item in [
        "Docker 실행 및 API/DB 조회 절차를 README에 고정하여 재현성을 높인다.",
        "통합 시나리오를 자동화 테스트로 확장하여 수동 캡처 의존도를 줄인다.",
        "프론트엔드 테스트 코드와 실제 컴포넌트 구조를 동기화하여 단위테스트 실패 항목을 보완한다.",
        "FastAPI to Spring Boot 저장 실패 시 재시도 횟수, 오류 메시지, 실패 payload를 별도 로그로 남기는 방식을 강화한다.",
        "프론트엔드에서 API 장애 상태를 사용자에게 명확히 표시할 수 있도록 error/empty/loading 상태 검증을 통합 시나리오에 추가한다.",
    ]:
        para(doc, f"- {item}")
    doc.save(OUT / "[KDT] 인천부평_1기_1조_통합테스트_결과서_양식맞춤_최종.docx")


def main():
    OUT.mkdir(exist_ok=True)
    write_md()
    build_plan()
    build_result()
    print(OUT)


if __name__ == "__main__":
    main()
