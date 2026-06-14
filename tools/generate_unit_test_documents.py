from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from PIL import Image


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "산출물_단위테스트"
PROJECT_NAME = "번호판 인식 스마트 교통안전 시스템"
TEAM_NAME = "인천부평_1기_1조"
ORG_NAME = "에스트래픽 스마트 모빌리티 DX Academy"
DOC_DATE = "2026. 06. 04"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_korean_font(run, size: int | None = None, bold: bool | None = None) -> None:
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def add_paragraph(doc: Document, text: str = "", style: str | None = None):
    p = doc.add_paragraph(style=style)
    r = p.add_run(text)
    set_korean_font(r, 10)
    return p


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(level=level)
    r = p.add_run(text)
    set_korean_font(r, 14 if level == 1 else 12, True)
    return p


def add_cover(doc: Document, title: str, subtitle: str) -> None:
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    set_korean_font(r, 22, True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"<{PROJECT_NAME}>")
    set_korean_font(r, 16, True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    set_korean_font(r, 12)

    for _ in range(7):
        doc.add_paragraph()

    for text in [DOC_DATE, ORG_NAME, TEAM_NAME]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_korean_font(r, 12, text == TEAM_NAME)

    doc.add_page_break()


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr[idx].text = header
        set_cell_shading(hdr[idx], "D9EAF7")
        for p in hdr[idx].paragraphs:
            for r in p.runs:
                set_korean_font(r, 9, True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
            for p in cells[idx].paragraphs:
                for r in p.runs:
                    set_korean_font(r, 8)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)
    doc.add_paragraph()
    return table


def add_image(doc: Document, image_name: str, caption: str) -> None:
    image_path = ROOT / image_name
    if not image_path.exists():
        add_paragraph(doc, f"[증빙 누락] {image_name} - {caption}")
        return

    add_paragraph(doc, f"증빙: {caption} ({image_name})")
    with Image.open(image_path) as im:
        w, h = im.size
    max_width_cm = 15.8
    max_height_cm = 10.5
    width_cm = max_width_cm
    height_cm = width_cm * h / w
    if height_cm > max_height_cm:
        height_cm = max_height_cm
        width_cm = height_cm * w / h
    doc.add_picture(str(image_path), width=Cm(width_cm))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def setup_doc() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    styles = doc.styles
    styles["Normal"].font.name = "Malgun Gothic"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    styles["Normal"].font.size = Pt(10)
    return doc


PLAN_MD = f"""# 단위테스트 계획서

## 문서 정보

- 프로젝트명: {PROJECT_NAME}
- 팀명: {TEAM_NAME}
- 작성일: {DOC_DATE}
- 작성 기준: Docker Compose 기반 최종 실행본, FastAPI/Spring Boot/Vue 3/PostgreSQL 연동 구조

## 1. 테스트 계획

### 1.1 테스트 목적

- FastAPI, Spring Boot, Vue 프론트엔드 주요 기능의 정상 동작 여부를 검증한다.
- API 계약, 응답 형식, DB 저장 결과, 화면 표시 상태가 최종 실행본 기준으로 일치하는지 확인한다.
- 자동화 테스트와 수동 API/DB 확인을 병행하여 발표 및 시연 환경에서 발생 가능한 결함을 조기에 식별한다.
- 테스트 수행 결과와 캡처 증빙을 남겨 산출물 제출 시 재현 가능한 근거를 확보한다.

### 1.2 테스트 범위

- FastAPI: Health Check, 이미지/영상 감지 처리, 번호판 정규화, 속도 설정/측정 로직, Spring Boot 전송 클라이언트
- Spring Boot: 감지 로그 저장, 분석 결과 저장, 중복 감지 처리, OCR 실패 처리, 과속 위반 조회/상태 처리
- Vue 프론트엔드: 주요 화면 렌더링, 감지 로그/대시보드 표시, 상태 컴포넌트 및 교통 혼잡도 유틸리티
- PostgreSQL: detection_logs, detection_analysis_results, speed_violations 저장 및 조회
- Docker Compose: PostgreSQL, Spring Boot, FastAPI, Frontend 컨테이너 빌드 및 실행

### 1.3 테스트 제외 범위

- YOLO/PaddleOCR 모델 자체의 학습 성능 평가
- 대규모 부하 테스트 및 장시간 안정성 테스트
- 운영 보안 진단, 침투 테스트, 접근 제어 정책 검증

### 1.4 테스트 일정

| 단계 | 일정 | 담당 | 비고 |
|---|---|---|---|
| 테스트 환경 준비 | 2026-06-04 | QA/통합 담당 | Docker Compose 및 의존성 설치 확인 |
| 단위테스트 실행 | 2026-06-04 | 각 파트 담당/QA | pytest, Gradle test, Vitest, Vite build |
| API/DB 확인 | 2026-06-04 | QA/통합 담당 | curl/Swagger/psql 기반 증빙 수집 |
| 결과서 작성 | 2026-06-04 | QA/통합 담당 | 캡처 이미지 및 결과 표 정리 |

### 1.5 역할 및 책임

| 역할 | 책임 |
|---|---|
| PM | 테스트 범위 승인 및 최종 산출물 검토 |
| FastAPI 담당 | AI 분석 API, OCR/번호판 처리, Spring Boot 전송 기능 검증 |
| Spring Boot 담당 | 저장 API, 조회 API, DB 연동, 예외 처리 검증 |
| Frontend 담당 | 화면 렌더링, API 응답 표시, 상태 처리 검증 |
| QA/통합 담당 | Docker 실행 환경, 전체 테스트 수행, 캡처 증빙 수집, 결과서 작성 |

## 2. 테스트 전략

### 2.1 테스트 유형

- 자동화 단위테스트: pytest, JUnit/Gradle, Vitest
- 빌드 검증: Vite production build
- API Smoke Test: FastAPI 및 Spring Boot 주요 API 응답 확인
- DB 저장 확인: PostgreSQL 테이블 조회
- 화면 확인: 프론트엔드 주요 화면 및 감지 결과 표시 확인

### 2.2 테스트 환경

| 구분 | 내용 |
|---|---|
| OS | Windows |
| Container | Docker Compose |
| DB | PostgreSQL 16 |
| Backend | Spring Boot 3.5.14, Java 21 |
| AI API | FastAPI, Python 3.12 |
| Frontend | Vue 3, Vite 5 |
| Node/NPM | Node v24.13.1, npm 11.8.0 |

### 2.3 테스트 도구

- pytest
- Gradle/JUnit
- Vitest
- Vite build
- Docker Compose
- Swagger UI
- curl/PowerShell Invoke-RestMethod
- PostgreSQL psql

## 3. 테스트 실행 계획

| 테스트 ID | 구분 | 테스트 항목 | 검증 기준 |
|---|---|---|---|
| UT-FAST-001 | FastAPI | pytest 자동화 테스트 | 전체 테스트 통과, 기능 실패 없음 |
| UT-FAST-002 | FastAPI | Health Check API | `/health` 200 응답 및 service/status 확인 |
| UT-SPR-001 | Spring Boot | Gradle/JUnit 테스트 | Spring context 및 감지 로그 통합 테스트 통과 |
| UT-FE-001 | Frontend | Vitest 테스트 | 컴포넌트/컴포저블 테스트 수행 결과 확인 |
| UT-FE-002 | Frontend | Vite production build | `npm run build` 정상 완료 |
| UT-API-001 | API | 감지 로그 조회 API | `/api/v1/detection-logs` 응답 확인 |
| UT-API-002 | API | 과속 위반 조회 API | `/api/speed-violations` 응답 확인 |
| UT-DB-001 | DB | 감지 로그 저장 조회 | `detection_logs` 최신 데이터 조회 가능 |
| UT-DB-002 | DB | 분석 결과 저장 조회 | `detection_analysis_results` 최신 데이터 조회 가능 |
| UT-DB-003 | DB | 과속 위반 조회 | `speed_violations` 조회 가능 |
| UT-INT-001 | 연동 | FastAPI to Spring Boot 저장 | FastAPI 분석 결과가 Spring Boot 저장 API로 전달됨 |
| UT-INT-002 | 화면 | 프론트 감지 결과 표시 | 프론트 화면에서 감지/조회 결과 확인 가능 |

## 4. 성공/실패 기준

- 성공: 명령어 종료 코드 0, API 200 응답, DB 조회 성공, 화면 표시 정상
- 실패: 테스트 suite 실패, API 4xx/5xx 오류, DB 저장 누락, 화면 표시 불가
- 경고: 테스트는 통과했으나 deprecation, chunk size, 캐시 도구 미설치 등 개선 권고가 있는 경우
"""


RESULT_MD = f"""# 단위테스트 결과서

## 문서 정보

- 프로젝트명: {PROJECT_NAME}
- 팀명: {TEAM_NAME}
- 작성일: {DOC_DATE}
- 작성 기준: 최종 Docker Compose 실행본 및 캡처 증빙

## 1. 테스트 결과 요약

| 구분 | 수행 결과 |
|---|---|
| FastAPI 자동화 테스트 | 64건 통과, 13건 warning 발생 |
| Spring Boot 테스트 | Gradle test BUILD SUCCESSFUL |
| Frontend Vitest | 1개 suite 실패, 3개 테스트 통과 |
| Frontend Build | Vite production build 정상 완료 |
| API 응답 확인 | FastAPI Health, Detection Logs, Speed Violations 응답 확인 |
| DB 저장 확인 | detection_logs, detection_analysis_results, speed_violations 조회 확인 |
| 연동 확인 | FastAPI 분석 결과의 Spring Boot/API 저장 흐름 확인 |

## 2. 상세 결과

| 테스트 ID | 항목 | 수행 내용 | 결과 | 비고 |
|---|---|---|---|---|
| UT-FAST-001 | FastAPI pytest | `python -m pytest` 실행 | 정상 | 64 passed, 13 warnings |
| UT-FAST-002 | FastAPI Health | `/health` 응답 확인 | 정상 | status/service 응답 확인 |
| UT-SPR-001 | Spring Boot Gradle test | `gradlew.bat test` 실행 | 정상 | BUILD SUCCESSFUL |
| UT-FE-001 | Frontend Vitest | `npm test` 실행 | 오류 | `DataState.vue` import 경로/파일 누락으로 suite 실패 |
| UT-FE-002 | Frontend Build | `npm run build` 실행 | 정상 | Vite build 완료, chunk size warning 발생 |
| UT-API-001 | Detection Logs API | `/api/v1/detection-logs` 조회 | 정상 | API 응답 캡처 확인 |
| UT-API-002 | Speed Violations API | `/api/speed-violations` 조회 | 정상 | API 응답 캡처 확인 |
| UT-DB-001 | detection_logs DB | 최신 감지 로그 조회 | 정상 | PostgreSQL psql 조회 성공 |
| UT-DB-002 | detection_analysis_results DB | 최신 분석 결과 조회 | 정상 | PostgreSQL psql 조회 성공 |
| UT-DB-003 | speed_violations DB | 과속 위반 테이블 조회 | 정상 | PostgreSQL psql 조회 성공 |
| UT-INT-001 | FastAPI-Spring 연동 | FastAPI 분석 결과 Spring Boot 저장 | 정상 | 저장 API 응답 확인 |
| UT-INT-002 | 프론트 화면 | 메인/감지 로그/분석 화면 표시 | 정상 | 화면 캡처 확인 |

## 3. 경고 및 결함

### 3.1 FastAPI warning

- Paddle 확장 모듈에서 `ccache` 미설치 경고가 발생하였다.
- PaddleOCR `ocr()` 호출에 대해 deprecation warning이 발생하였다.
- 두 항목 모두 테스트 실패가 아니며 기능 오류는 발생하지 않았다.
- 향후 `predict()` 기반 호출 방식 전환과 ccache 설치를 검토한다.

### 3.2 Frontend Vitest 오류

- `tests/DataState.test.js`에서 `@/components/dashboard/DataState.vue` import를 수행하지만 해당 파일이 현재 소스 트리에 존재하지 않아 suite가 실패하였다.
- `useOSMRoads.test.js`의 혼잡도 유틸리티 테스트 3건은 통과하였다.
- 조치 방향: `DataState.vue` 컴포넌트를 복구하거나 테스트 대상을 현재 컴포넌트 구조에 맞게 수정한다.

### 3.3 Frontend build warning

- Vite CJS Node API deprecation warning과 일부 chunk size warning이 발생하였다.
- production build는 정상 완료되었으며, 기능 실패는 아니다.
- 향후 dynamic import 또는 manualChunks 적용을 검토한다.

## 4. 개선사항

- 프론트엔드 테스트 코드와 실제 컴포넌트 구조를 동기화한다.
- PaddleOCR 호출부를 최신 API 권고 방식으로 변경한다.
- Docker 기준 실행 절차, API 확인 명령어, DB 조회 쿼리를 README 또는 운영 문서에 고정한다.
- 단위테스트와 통합테스트를 분리하여 테스트 ID, API 계약, DB 테이블 증빙을 지속 관리한다.
"""


RESULT_IMAGES = [
    ("UT_FAST_001_pytest_result.png", "FastAPI pytest 64건 통과 결과"),
    ("UT_SPR_001_gradle_test_result.png", "Spring Boot Gradle test BUILD SUCCESSFUL"),
    ("UT_FE_001_npm_test_result.png", "Frontend Vitest 수행 결과"),
    ("API_001_fastapi_health.png", "FastAPI Health Check API 응답"),
    ("API_002_detection_logs_response.png", "Spring Boot 감지 로그 조회 API 응답"),
    ("API_003_speed_violations_response.png", "Spring Boot 과속 위반 조회 API 응답"),
    ("IT_001_docker_compose_build.png", "Docker Compose 빌드 및 실행 로그"),
    ("IT_003_front_to_fastapi_detection.png", "프론트엔드에서 FastAPI 분석 결과 확인"),
    ("IT_004_fastapi_to_spring_saved.png", "FastAPI 분석 결과의 Spring Boot 저장 응답"),
    ("IT_005_01_detection_logs_query.png", "PostgreSQL detection_logs 조회 결과"),
    ("IT_005_02_detection_analysis_results_query.png", "PostgreSQL detection_analysis_results 조회 결과"),
    ("IT_005_03_speed_violations_query.png", "PostgreSQL speed_violations 조회 결과"),
    ("FE_001_main_page.png", "프론트엔드 메인 화면"),
    ("FE_003_detection_logs.png", "프론트엔드 감지 로그 화면"),
]


def write_markdown(path: Path, text: str) -> None:
    path.write_text(text.strip() + "\n", encoding="utf-8")


def build_plan_docx(path: Path) -> None:
    doc = setup_doc()
    add_cover(doc, "단위테스트 계획서", "최종본")

    add_heading(doc, "1. 테스트 계획")
    for title, body in [
        ("1.1 테스트 목적", [
            "FastAPI, Spring Boot, Vue 프론트엔드 주요 기능의 정상 동작 여부를 검증한다.",
            "API 계약, 응답 형식, DB 저장 결과, 화면 표시 상태가 최종 실행본 기준으로 일치하는지 확인한다.",
            "자동화 테스트와 수동 API/DB 확인을 병행하여 결함을 조기에 식별한다.",
        ]),
        ("1.2 테스트 범위", [
            "FastAPI: Health Check, 이미지/영상 감지 처리, 번호판 정규화, 속도 설정/측정 로직, Spring Boot 전송 클라이언트",
            "Spring Boot: 감지 로그 저장, 분석 결과 저장, 중복 감지 처리, OCR 실패 처리, 과속 위반 조회/상태 처리",
            "Vue 프론트엔드: 주요 화면 렌더링, 감지 로그/대시보드 표시, 상태 컴포넌트 및 교통 혼잡도 유틸리티",
            "PostgreSQL: detection_logs, detection_analysis_results, speed_violations 저장 및 조회",
            "Docker Compose: PostgreSQL, Spring Boot, FastAPI, Frontend 컨테이너 빌드 및 실행",
        ]),
        ("1.3 테스트 제외 범위", [
            "YOLO/PaddleOCR 모델 자체의 학습 성능 평가",
            "대규모 부하 테스트 및 장시간 안정성 테스트",
            "운영 보안 진단, 침투 테스트, 접근 제어 정책 검증",
        ]),
    ]:
        add_heading(doc, title, 2)
        for item in body:
            add_paragraph(doc, f"- {item}")

    add_heading(doc, "1.4 테스트 일정", 2)
    add_table(
        doc,
        ["단계", "일정", "담당", "비고"],
        [
            ["테스트 환경 준비", "2026-06-04", "QA/통합 담당", "Docker Compose 및 의존성 설치 확인"],
            ["단위테스트 실행", "2026-06-04", "각 파트 담당/QA", "pytest, Gradle test, Vitest, Vite build"],
            ["API/DB 확인", "2026-06-04", "QA/통합 담당", "curl/Swagger/psql 기반 증빙 수집"],
            ["결과서 작성", "2026-06-04", "QA/통합 담당", "캡처 이미지 및 결과 표 정리"],
        ],
        [3.5, 3.5, 3, 6],
    )

    add_heading(doc, "1.5 역할 및 책임", 2)
    add_table(
        doc,
        ["역할", "책임"],
        [
            ["PM", "테스트 범위 승인 및 최종 산출물 검토"],
            ["FastAPI 담당", "AI 분석 API, OCR/번호판 처리, Spring Boot 전송 기능 검증"],
            ["Spring Boot 담당", "저장 API, 조회 API, DB 연동, 예외 처리 검증"],
            ["Frontend 담당", "화면 렌더링, API 응답 표시, 상태 처리 검증"],
            ["QA/통합 담당", "Docker 실행 환경, 전체 테스트 수행, 캡처 증빙 수집, 결과서 작성"],
        ],
        [4, 12],
    )

    add_heading(doc, "2. 테스트 전략")
    add_heading(doc, "2.1 테스트 유형", 2)
    for item in ["자동화 단위테스트", "빌드 검증", "API Smoke Test", "DB 저장 확인", "화면 확인"]:
        add_paragraph(doc, f"- {item}")

    add_heading(doc, "2.2 테스트 환경", 2)
    add_table(
        doc,
        ["구분", "내용"],
        [
            ["OS", "Windows"],
            ["Container", "Docker Compose"],
            ["DB", "PostgreSQL 16"],
            ["Backend", "Spring Boot 3.5.14, Java 21"],
            ["AI API", "FastAPI, Python 3.12"],
            ["Frontend", "Vue 3, Vite 5"],
            ["Node/NPM", "Node v24.13.1, npm 11.8.0"],
        ],
        [4, 12],
    )

    add_heading(doc, "2.3 테스트 도구", 2)
    add_paragraph(doc, "pytest, Gradle/JUnit, Vitest, Vite build, Docker Compose, Swagger UI, curl/PowerShell Invoke-RestMethod, PostgreSQL psql")

    add_heading(doc, "3. 테스트 실행 계획")
    add_table(
        doc,
        ["테스트 ID", "구분", "테스트 항목", "검증 기준"],
        [
            ["UT-FAST-001", "FastAPI", "pytest 자동화 테스트", "전체 테스트 통과, 기능 실패 없음"],
            ["UT-FAST-002", "FastAPI", "Health Check API", "/health 200 응답 및 service/status 확인"],
            ["UT-SPR-001", "Spring Boot", "Gradle/JUnit 테스트", "Spring context 및 감지 로그 통합 테스트 통과"],
            ["UT-FE-001", "Frontend", "Vitest 테스트", "컴포넌트/컴포저블 테스트 수행 결과 확인"],
            ["UT-FE-002", "Frontend", "Vite production build", "npm run build 정상 완료"],
            ["UT-API-001", "API", "감지 로그 조회 API", "/api/v1/detection-logs 응답 확인"],
            ["UT-API-002", "API", "과속 위반 조회 API", "/api/speed-violations 응답 확인"],
            ["UT-DB-001", "DB", "감지 로그 저장 조회", "detection_logs 최신 데이터 조회 가능"],
            ["UT-DB-002", "DB", "분석 결과 저장 조회", "detection_analysis_results 최신 데이터 조회 가능"],
            ["UT-DB-003", "DB", "과속 위반 조회", "speed_violations 조회 가능"],
            ["UT-INT-001", "연동", "FastAPI to Spring Boot 저장", "FastAPI 분석 결과가 Spring Boot 저장 API로 전달됨"],
            ["UT-INT-002", "화면", "프론트 감지 결과 표시", "프론트 화면에서 감지/조회 결과 확인 가능"],
        ],
        [3, 3, 5, 6],
    )

    add_heading(doc, "4. 성공/실패 기준")
    for item in [
        "성공: 명령어 종료 코드 0, API 200 응답, DB 조회 성공, 화면 표시 정상",
        "실패: 테스트 suite 실패, API 4xx/5xx 오류, DB 저장 누락, 화면 표시 불가",
        "경고: 테스트는 통과했으나 deprecation, chunk size, 캐시 도구 미설치 등 개선 권고가 있는 경우",
    ]:
        add_paragraph(doc, f"- {item}")

    doc.save(path)


def build_result_docx(path: Path) -> None:
    doc = setup_doc()
    add_cover(doc, "단위테스트 결과서", "최종본")

    add_heading(doc, "1. 테스트 결과 요약")
    add_table(
        doc,
        ["구분", "수행 결과"],
        [
            ["FastAPI 자동화 테스트", "64건 통과, 13건 warning 발생"],
            ["Spring Boot 테스트", "Gradle test BUILD SUCCESSFUL"],
            ["Frontend Vitest", "1개 suite 실패, 3개 테스트 통과"],
            ["Frontend Build", "Vite production build 정상 완료"],
            ["API 응답 확인", "FastAPI Health, Detection Logs, Speed Violations 응답 확인"],
            ["DB 저장 확인", "detection_logs, detection_analysis_results, speed_violations 조회 확인"],
            ["연동 확인", "FastAPI 분석 결과의 Spring Boot/API 저장 흐름 확인"],
        ],
        [5, 11],
    )

    add_heading(doc, "2. 상세 결과")
    add_table(
        doc,
        ["테스트 ID", "항목", "수행 내용", "결과", "비고"],
        [
            ["UT-FAST-001", "FastAPI pytest", "python -m pytest 실행", "정상", "64 passed, 13 warnings"],
            ["UT-FAST-002", "FastAPI Health", "/health 응답 확인", "정상", "status/service 응답 확인"],
            ["UT-SPR-001", "Spring Boot Gradle test", "gradlew.bat test 실행", "정상", "BUILD SUCCESSFUL"],
            ["UT-FE-001", "Frontend Vitest", "npm test 실행", "오류", "DataState.vue import 경로/파일 누락"],
            ["UT-FE-002", "Frontend Build", "npm run build 실행", "정상", "Vite build 완료, chunk size warning 발생"],
            ["UT-API-001", "Detection Logs API", "/api/v1/detection-logs 조회", "정상", "API 응답 캡처 확인"],
            ["UT-API-002", "Speed Violations API", "/api/speed-violations 조회", "정상", "API 응답 캡처 확인"],
            ["UT-DB-001", "detection_logs DB", "최신 감지 로그 조회", "정상", "PostgreSQL psql 조회 성공"],
            ["UT-DB-002", "detection_analysis_results DB", "최신 분석 결과 조회", "정상", "PostgreSQL psql 조회 성공"],
            ["UT-DB-003", "speed_violations DB", "과속 위반 테이블 조회", "정상", "PostgreSQL psql 조회 성공"],
            ["UT-INT-001", "FastAPI-Spring 연동", "FastAPI 분석 결과 Spring Boot 저장", "정상", "저장 API 응답 확인"],
            ["UT-INT-002", "프론트 화면", "메인/감지 로그/분석 화면 표시", "정상", "화면 캡처 확인"],
        ],
        [2.7, 3.8, 4.8, 2, 4.7],
    )

    add_heading(doc, "3. 경고 및 결함")
    add_heading(doc, "3.1 FastAPI warning", 2)
    for item in [
        "Paddle 확장 모듈에서 ccache 미설치 경고가 발생하였다.",
        "PaddleOCR ocr() 호출에 대해 deprecation warning이 발생하였다.",
        "두 항목 모두 테스트 실패가 아니며 기능 오류는 발생하지 않았다.",
        "향후 predict() 기반 호출 방식 전환과 ccache 설치를 검토한다.",
    ]:
        add_paragraph(doc, f"- {item}")

    add_heading(doc, "3.2 Frontend Vitest 오류", 2)
    for item in [
        "tests/DataState.test.js에서 @/components/dashboard/DataState.vue import를 수행하지만 해당 파일이 현재 소스 트리에 존재하지 않아 suite가 실패하였다.",
        "useOSMRoads.test.js의 혼잡도 유틸리티 테스트 3건은 통과하였다.",
        "조치 방향: DataState.vue 컴포넌트를 복구하거나 테스트 대상을 현재 컴포넌트 구조에 맞게 수정한다.",
    ]:
        add_paragraph(doc, f"- {item}")

    add_heading(doc, "3.3 Frontend build warning", 2)
    for item in [
        "Vite CJS Node API deprecation warning과 일부 chunk size warning이 발생하였다.",
        "production build는 정상 완료되었으며, 기능 실패는 아니다.",
        "향후 dynamic import 또는 manualChunks 적용을 검토한다.",
    ]:
        add_paragraph(doc, f"- {item}")

    add_heading(doc, "4. 개선사항")
    for item in [
        "프론트엔드 테스트 코드와 실제 컴포넌트 구조를 동기화한다.",
        "PaddleOCR 호출부를 최신 API 권고 방식으로 변경한다.",
        "Docker 기준 실행 절차, API 확인 명령어, DB 조회 쿼리를 README 또는 운영 문서에 고정한다.",
        "단위테스트와 통합테스트를 분리하여 테스트 ID, API 계약, DB 테이블 증빙을 지속 관리한다.",
    ]:
        add_paragraph(doc, f"- {item}")

    add_heading(doc, "5. 증빙 이미지")
    for image_name, caption in RESULT_IMAGES:
        add_heading(doc, caption, 2)
        add_image(doc, image_name, caption)

    doc.save(path)


def main() -> None:
    OUT.mkdir(exist_ok=True)
    write_markdown(OUT / "[KDT] 인천부평_1기_1조_단위테스트_계획서_최종.md", PLAN_MD)
    write_markdown(OUT / "[KDT] 인천부평_1기_1조_단위테스트_결과서_최종.md", RESULT_MD)
    build_plan_docx(OUT / "[KDT] 인천부평_1기_1조_단위테스트_계획서_최종.docx")
    build_result_docx(OUT / "[KDT] 인천부평_1기_1조_단위테스트_결과서_최종.docx")
    print(OUT)


if __name__ == "__main__":
    main()
